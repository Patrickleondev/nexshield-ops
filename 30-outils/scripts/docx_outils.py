"""Briques de construction DOCX partagées, à la charte.

Utilisé par generer_documents.py (livrables) et generer_juridique.py (contrats).
Toute mise en forme passe par ici : c'est ce qui garantit qu'un contrat et un
rapport se ressemblent.
"""
from __future__ import annotations
import os, sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import charte as C

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

STYLE_PARAGRAPHE = 1   # WD_STYLE_TYPE.PARAGRAPH


def fond(element, hexa: str) -> None:
    ombre = OxmlElement("w:shd")
    ombre.set(qn("w:val"), "clear")
    ombre.set(qn("w:fill"), hexa)
    element.append(ombre)


def definir_style(doc, nom, taille, gras, couleur, avant=0, apres=6, police=None):
    existants = [s.name for s in doc.styles]
    st = doc.styles[nom] if nom in existants else doc.styles.add_style(nom, STYLE_PARAGRAPHE)
    st.font.name = police or C.POLICE_CORPS
    st.font.size = Pt(taille)
    st.font.bold = gras
    st.font.color.rgb = RGBColor(*C.rgb(couleur))
    st.paragraph_format.space_before = Pt(avant)
    st.paragraph_format.space_after = Pt(apres)
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:ascii"), police or C.POLICE_CORPS)
    rf.set(qn("w:hAnsi"), police or C.POLICE_CORPS)
    rf.set(qn("w:cs"), C.REPLI_SANS)
    return st


def nouveau_document(marge_cm: float = 2.5) -> Document:
    """Document vierge avec la charte appliquée."""
    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(marge_cm)
        s.left_margin = s.right_margin = Cm(marge_cm)

    n = doc.styles["Normal"]
    n.font.name = C.POLICE_CORPS
    n.font.size = Pt(C.TAILLES["corps"])
    n.font.color.rgb = RGBColor(*C.rgb(C.TEXTE))
    n.paragraph_format.space_after = Pt(6)
    n.paragraph_format.line_spacing = 1.25

    definir_style(doc, "Heading 1", C.TAILLES["t1"], True, C.PRIMAIRE, 18, 8, C.POLICE_TITRE)
    definir_style(doc, "Heading 2", C.TAILLES["t2"], True, C.SECONDAIRE, 14, 6, C.POLICE_TITRE)
    definir_style(doc, "Heading 3", C.TAILLES["t3"], True, C.SECONDAIRE, 10, 4, C.POLICE_TITRE)
    definir_style(doc, "NS Legende", C.TAILLES["petit"], False, C.TEXTE_FAIBLE, 0, 10)
    definir_style(doc, "NS Code", C.TAILLES["mono"], False, C.TEXTE, 4, 8, C.POLICE_MONO)
    definir_style(doc, "NS Encadre", C.TAILLES["corps"], False, C.SECONDAIRE, 6, 10)
    definir_style(doc, "NS Clause", C.TAILLES["corps"], False, C.TEXTE, 4, 8)
    return doc


def pied_de_page(section, texte_gauche: str) -> None:
    p = section.footer.paragraphs[0]
    p.text = ""
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(f"{texte_gauche}\t")
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(*C.rgb(C.TEXTE_FAIBLE))
    r.font.name = C.POLICE_CORPS
    for instr in ("PAGE", "NUMPAGES"):
        if instr == "NUMPAGES":
            s = p.add_run(" / "); s.font.size = Pt(8)
            s.font.color.rgb = RGBColor(*C.rgb(C.TEXTE_FAIBLE))
        deb = OxmlElement("w:fldChar"); deb.set(qn("w:fldCharType"), "begin")
        txt = OxmlElement("w:instrText"); txt.set(qn("xml:space"), "preserve")
        txt.text = f" {instr} "
        fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
        run = p.add_run(); run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(*C.rgb(C.TEXTE_FAIBLE))
        run._r.append(deb); run._r.append(txt); run._r.append(fin)


def tableau(doc, entetes: list[str], lignes: list[list[str]], largeurs=None,
            entete_visible: bool = True):
    t = doc.add_table(rows=1, cols=len(entetes))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, e in enumerate(entetes):
        cel = t.rows[0].cells[i]
        cel.text = ""
        r = cel.paragraphs[0].add_run(e)
        r.font.bold = True
        r.font.size = Pt(9.5)
        r.font.name = C.POLICE_TITRE
        r.font.color.rgb = RGBColor(*C.rgb(C.BLANC if entete_visible else C.TEXTE))
        if entete_visible:
            fond(cel._tc.get_or_add_tcPr(), C.PRIMAIRE)
    for n, ligne in enumerate(lignes):
        cells = t.add_row().cells
        for i, v in enumerate(ligne):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(v)
            r.font.size = Pt(9.5)
            r.font.name = C.POLICE_CORPS
            r.font.color.rgb = RGBColor(*C.rgb(C.TEXTE))
            if n % 2 == 1:
                fond(cells[i]._tc.get_or_add_tcPr(), C.FOND_ALTERNE)
    if largeurs:
        for i, l in enumerate(largeurs):
            for row in t.rows:
                row.cells[i].width = Cm(l)
    doc.add_paragraph()
    return t


def sommaire(doc, titre: str = "Sommaire", profondeur: str = "1-3") -> None:
    doc.add_heading(titre, level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    deb = OxmlElement("w:fldChar"); deb.set(qn("w:fldCharType"), "begin")
    txt = OxmlElement("w:instrText"); txt.set(qn("xml:space"), "preserve")
    txt.text = f' TOC \\o "{profondeur}" \\h \\z \\u '
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    ph = OxmlElement("w:t"); ph.text = "Clic droit puis « Mettre à jour les champs » (F9)"
    fin = OxmlElement("w:fldChar"); fin.set(qn("w:fldCharType"), "end")
    for e in (deb, txt, sep, ph, fin):
        run._r.append(e)


def couverture(doc, titre: str, sous_titre: str, metadonnees: list[list[str]],
               mention: str | None = None) -> None:
    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(C.SOCIETE)
    r.font.size, r.font.bold, r.font.name = Pt(34), True, C.POLICE_TITRE
    r.font.color.rgb = RGBColor(*C.rgb(C.PRIMAIRE))
    p = doc.add_paragraph()
    r = p.add_run(C.BASELINE)
    r.font.size, r.font.name = Pt(11), C.POLICE_CORPS
    r.font.color.rgb = RGBColor(*C.rgb(C.ACCENT))

    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run(titre)
    r.font.size, r.font.bold, r.font.name = Pt(24), True, C.POLICE_TITRE
    r.font.color.rgb = RGBColor(*C.rgb(C.SECONDAIRE))
    if sous_titre:
        p = doc.add_paragraph()
        r = p.add_run(sous_titre)
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(*C.rgb(C.TEXTE_FAIBLE))

    doc.add_paragraph()
    tableau(doc, ["", ""], metadonnees, largeurs=[5.5, 10.5], entete_visible=False)

    if mention:
        p = doc.add_paragraph(style="NS Legende")
        p.add_run(mention)


def consigne(doc, texte: str) -> None:
    """Consigne de rédaction interne, à supprimer avant remise au client."""
    p = doc.add_paragraph(style="NS Encadre")
    r = p.add_run("[Consigne interne — supprimer avant remise] " + texte)
    r.italic = True
    r.font.color.rgb = RGBColor(*C.rgb(C.ACCENT))


def bloc_signatures(doc, parties: list[str]) -> None:
    lignes = [["Nom et prénom"] + [""] * len(parties),
              ["Fonction"] + [""] * len(parties),
              ["Date"] + [""] * len(parties),
              ["Signature et cachet"] + [""] * len(parties)]
    largeur = 16 / (len(parties) + 1)
    tableau(doc, [""] + parties, lignes, largeurs=[largeur] * (len(parties) + 1))


def historique_versions(doc) -> None:
    doc.add_heading("Historique des versions", level=2)
    tableau(doc, ["Version", "Date", "Auteur", "Nature de la modification"],
            [["v0.1", "<AAAA-MM-JJ>", "<…>", "Création"],
             ["", "", "", ""]],
            largeurs=[2.2, 3, 4, 6.8])
